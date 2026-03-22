import base64
import io
import os
import tempfile
import uuid
import asyncio
import time
from PIL import Image
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from extensions import state
from history_rag import history_rag
from context_config import get_search_k
from llm_factory import create_llm


def get_embedding_model(base_url: str):
    return OllamaEmbeddings(
        model="nomic-embed-text",
        base_url=base_url
    )


def get_llm_model(temperature=0.7):
    provider = state.llm_provider
    
    if provider == "ollama":
        return create_llm(
            provider="ollama",
            model="qwen3.5:4b",
            base_url=state.ollama_base_url,
            temperature=temperature
        )
    elif provider == "openai":
        return create_llm(
            provider="openai",
            model=state.openai_current_model if hasattr(state, 'openai_current_model') and state.openai_current_model else model_name,
            base_url=state.get_openai_base_url() if hasattr(state, 'get_openai_base_url') else None,
            api_key=state.get_openai_api_key() if hasattr(state, 'get_openai_api_key') else None,
            temperature=temperature
        )
    elif provider == "anthropic":
        return create_llm(
            provider="anthropic",
            model=state.anthropic_current_model if hasattr(state, 'anthropic_current_model') and state.anthropic_current_model else "claude-3-sonnet-20240229",
            base_url=state.get_anthropic_base_url() if hasattr(state, 'get_anthropic_base_url') else None,
            api_key=state.get_anthropic_api_key() if hasattr(state, 'get_anthropic_api_key') else None,
            temperature=temperature
        )
    else:
        raise ValueError(f"Unknown provider: {provider}")


def load_document(file_path, file_type):
    if file_type == "pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif file_type == "docx":
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": file_path})]
    elif file_type == "txt":
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": file_path})]
    return []


def process_document(documents, base_url: str):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_documents(documents)
    
    if state.llm_provider == "ollama":
        embedding = get_embedding_model(base_url)
        vector_store = FAISS.from_documents(chunks, embedding)
        return vector_store
    else:
        return chunks


def generate_summary(messages):
    try:
        prompt = """请仔细阅读以下对话记录，然后生成一个简洁的摘要。

【对话记录】
"""
        for msg in messages:
            prompt += f"{msg.role}: {msg.content}\n"

        prompt += """
【要求】
1. 一句话概括对话核心主题
2. 如果有文档问答，记录文档主题
3. 记录用户的主要意图

【摘要格式】
主题：[核心主题]
文档：[文档名称或无]
意图：[用户主要目的]

请生成摘要："""
        
        llm = get_llm_model(temperature=0.3)
        response = llm.invoke(prompt)
        
        content = response.content
        if isinstance(content, list):
            text_parts = []
            for part in content:
                if isinstance(part, dict) and part.get('type') == 'text':
                    text_parts.append(part.get('text', ''))
                elif isinstance(part, str):
                    text_parts.append(part)
            content = ''.join(text_parts)
        
        return content.strip()
    except Exception as e:
        print(f"生成摘要失败：{str(e)}")
        return None


def process_image(image_data):
    try:
        if image_data.startswith('data:image'):
            image_data = image_data.split(',')[1]
        
        image_bytes = base64.b64decode(image_data)
        image = Image.open(io.BytesIO(image_bytes))
        
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        max_size = (512, 512)
        image.thumbnail(max_size, Image.LANCZOS)
        
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"temp_{uuid.uuid4()}.jpg")
        image.save(temp_path, 'JPEG', quality=70)
        
        return temp_path
    except Exception as e:
        print(f"图像处理失败: {str(e)}")
        return None


def encode_image_to_base64(image_path):
    try:
        with open(image_path, 'rb') as f:
            return base64.b64encode(f.read()).decode('utf-8')
    except Exception as e:
        print(f"图像编码失败: {str(e)}")
        return None


def prepare_messages(conversation, query, system_prompt, images=None):
    messages = [SystemMessage(content=system_prompt)]
    
    total_turns = conversation.get_total_turns()
    
    if total_turns > state.max_context_turns:
        if not conversation.summary:
            early_messages = conversation.messages[:state.max_context_turns * 2]
            summary = generate_summary(early_messages)
            if summary:
                conversation.summary = summary
        
        if conversation.summary:
            messages.append(SystemMessage(content=f"之前的对话摘要：{conversation.summary}"))
        
        recent_messages = conversation.messages[-state.max_context_turns * 2:]
        for msg in recent_messages:
            if msg.role == "user":
                messages.append(HumanMessage(content=msg.content))
            elif msg.role == "assistant":
                messages.append(AIMessage(content=msg.content))
    else:
        for msg in conversation.messages:
            if msg.role == "user":
                messages.append(HumanMessage(content=msg.content))
            elif msg.role == "assistant":
                messages.append(AIMessage(content=msg.content))
    
    if images and len(images) > 0:
        image_contents = []
        for img in images:
            image_url = f"data:image/jpeg;base64,{img['data']}"
            print(f"图片URL长度: {len(image_url)}")
            print(f"图片数据长度: {len(img['data'])}")
            image_contents.append({
                "type": "image_url",
                "image_url": {
                    "url": image_url
                }
            })
        
        content = [
            {"type": "text", "text": query}
        ] + image_contents
        
        messages.append(HumanMessage(content=content))
    else:
        messages.append(HumanMessage(content=query))
    
    return messages


async def generate_answer(query, model_name=None, mode="qa"):
    from agent import stream_graph
    from tools import news_toolkit
    from document_tools import document_tools, get_document_summary, get_document_outline
    
    try:
        conversation = state.get_current_conversation()
        
        if model_name is None:
            model_name = "qwen3.5:4b"
        
        print(f"开始生成回答，模型: {model_name}，模式: {mode} (LangGraph工作流)")
        
        conversation.add_message("user", query)
        
        images = conversation.images
        
        full_response = ""
        
        for chunk in stream_graph(query, model_name, images, mode):
            if state.should_stop:
                full_response += "\n\n操作已中断"
                state.response_queue.put(("chunk", "\n\n操作已中断"))
                break
            
            full_response += chunk
            state.response_queue.put(("chunk", chunk))
        
        if not state.should_stop:
            conversation.add_message("assistant", full_response)
            state.persist_message("user", query)
            state.persist_message("assistant", full_response)
            auto_name_conversation(conversation)
            state.response_queue.put(("done", ""))
        else:
            state.response_queue.put(("error", "操作已中断"))

    except Exception as e:
        print(f"生成回答失败: {str(e)}")
        import traceback
        traceback.print_exc()
        state.response_queue.put(("error", f"生成失败：{str(e)}"))
    finally:
        state.is_generating = False
        state.should_stop = False


def auto_name_conversation(conversation):
    if conversation.messages:
        user_assistant_pairs = 0
        for i in range(0, len(conversation.messages) - 1, 2):
            if (conversation.messages[i].role == "user" and
                i + 1 < len(conversation.messages) and
                conversation.messages[i + 1].role == "assistant"):
                user_assistant_pairs += 1

        if conversation.name == "新对话":
            for msg in conversation.messages:
                if msg.role == "user":
                    name = msg.content[:20] + ("..." if len(msg.content) > 20 else "")
                    conversation.name = name
                    state.persist_conversation_name(name)
                    break

        if user_assistant_pairs > 0 and user_assistant_pairs % 5 == 0:
            summary = generate_summary(conversation.messages)
            if summary:
                conversation.name = summary[:30] + ("..." if len(summary) > 30 else "")
                state.persist_conversation_name(conversation.name)
                state.persist_summary(summary)
