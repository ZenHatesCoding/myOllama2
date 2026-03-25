from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings


def get_embedding_model(base_url: str):
    return OllamaEmbeddings(
        model="nomic-embed-text",
        base_url=base_url
    )


def load_document(file_path, file_type):
    if file_type == "pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif file_type == "docx":
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": file_path})]
    elif file_type == "txt":
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": file_path})]
    return []


def process_document(documents, base_url: str):
    from core import state

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_documents(documents)
    
    if state.llm_provider == "ollama":
        embedding = get_embedding_model(base_url)
        vector_store = FAISS.from_documents(chunks, embedding)
        return vector_store
    else:
        return chunks
